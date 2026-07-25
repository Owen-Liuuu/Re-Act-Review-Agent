"""Generate a Word (.docx) evaluation report with comparison charts.

Produces a professional document with:
  - Cover page: review title, run metadata
  - Executive summary: overall verdict, key findings
  - Step 1-2: search & verification results
  - Step 3-4: per-paper data comparison table
      (Student vs LLM-A vs LLM-B, colour-coded)
  - Bar chart: agreement rates per paper
  - Bar chart: agreement rates per field
  - Appendix: all flags
"""
from __future__ import annotations

import io
import re
from dataclasses import dataclass, field
from datetime import datetime
from pathlib import Path

import structlog

from react_review.core.enums import FieldStatus
from react_review.pipeline.schemas import PipelineRunResult
from react_review.steps.table_comparison.schemas import (
    FieldDiff,
    TableComparisonResult,
)


# ----------------------------------------------------------------------
# Group-pair merging
# ----------------------------------------------------------------------
# The Step 4 comparator emits one FieldDiff per (canonical_base, group)
# pair, e.g. ``age_t1dm`` and ``age_control`` are two separate rows.
# But in the original "Characteristics of included studies" table the
# Age column is ONE column with two values per paper (a T1DM row and a
# Control row). To restore that layout in the report we merge group-pairs
# back into a single display row whose Student / AI cells contain both
# values labelled by group (e.g. "T1DM: 12.90±1.30\nControl: 12.96±1.12").
# Status is collapsed to the worst of the constituent statuses so a
# disagreement on either group is still surfaced.
# ----------------------------------------------------------------------

# Suffixes used by canonical_key to denote a group. Order matters only
# for display ordering inside the merged cell.
_GROUP_SUFFIXES: tuple[str, ...] = ("_t1dm", "_control")

# Pretty labels for the merged cell's per-group lines.
_GROUP_LABEL: dict[str, str] = {"_t1dm": "T1DM", "_control": "Control"}


def _strip_group_suffix(name: str) -> tuple[str, str]:
    """Split a canonical field name like ``age_t1dm`` into (``age``, ``_t1dm``).

    Returns ``(name, "")`` if no known group suffix is present.
    """
    lower = name.lower()
    for suf in _GROUP_SUFFIXES:
        if lower.endswith(suf):
            return name[: -len(suf)], suf
    return name, ""


# Status priority — higher = worse. Used to roll up grouped statuses
# so the merged row reflects the worst constituent.
_STATUS_PRIORITY: dict[str, int] = {
    "match": 0,
    "missing_student": 1,
    "missing_model": 1,
    "not_comparable": 1,
    "partial_match": 2,
    "needs_review": 3,
    "diff": 4,
}


@dataclass
class _RenderRow:
    """A single row to be rendered in the per-paper comparison table.

    May be either a passthrough of one FieldDiff or a synthesis of two
    grouped diffs (T1DM + Control collapsed into one row).

    Layout (post-2026-05-10 redesign):
      Field | Student | <LLM-A value> | <LLM-B value> | Status |
      Reason (LLM-A) | Reason (LLM-B)

    Each extractor gets its own ``Reason`` cell rather than sharing a
    single Evidence column. When an extractor failed (no value), the
    cell shows ``"MODEL N/A"``; otherwise it shows the supporting quote
    prefixed with ``[llm:<short_name>]``.
    """

    field_label: str
    student_text: str
    model_texts: list[str]               # value cells, one per extractor
    status_value: str                    # FieldStatus.value, e.g. "match"
    model_reasons: list[str]             # reason cells, one per extractor


def _short_extractor_name(eid: str) -> str:
    """Build a short display name for an extractor id.

    "qwen-plus-extractor"   → "Qwen-Plus"
    "kimi-k2.5-extractor"    → "Kimi-K2.5"
    "gemini-2.5-flash-extractor" → "Gemini-2.5"
    """
    parts = eid.replace("-extractor", "").split("-")
    return "-".join(p.capitalize() for p in parts[:2])


def _merge_grouped_diffs(
    diffs: list[FieldDiff],
    extractor_ids: list[str],
) -> list[_RenderRow]:
    """Collapse (base, t1dm)+(base, control) FieldDiffs into single rows.

    Diffs that don't have a recognised group suffix pass through unchanged.
    A diff with a group suffix is grouped with any sibling diffs that share
    the same base name; the resulting row labels each value with its group.
    """
    # Group by base name. We preserve first-seen order so the report's row
    # order matches the comparator's output order.
    base_to_indices: dict[str, list[int]] = {}
    base_order: list[str] = []
    for idx, d in enumerate(diffs):
        base, suf = _strip_group_suffix(d.field_name)
        # Only merge when at least one suffix is present; bare names go
        # through with their own base key so they remain solo rows.
        merge_key = base if suf else f"__solo__:{idx}"
        if merge_key not in base_to_indices:
            base_to_indices[merge_key] = []
            base_order.append(merge_key)
        base_to_indices[merge_key].append(idx)

    rows: list[_RenderRow] = []
    for key in base_order:
        members = [diffs[i] for i in base_to_indices[key]]
        if len(members) == 1 and not key.startswith("__solo__:"):
            # Solo grouped diff — only T1DM or only Control was emitted.
            # Keep the suffix visible so the reader sees which group it is.
            rows.append(_render_single(members[0], extractor_ids))
        elif len(members) == 1:
            rows.append(_render_single(members[0], extractor_ids))
        else:
            rows.append(_render_merged(key, members, extractor_ids))
    return rows


def _build_field_label(student_label: str, canonical: str) -> str:
    """Render the Field cell as up to two lines.

    Line 1 — student's verbatim column header (the join key).
    Line 2 — canonical concept the AI extracted under, only when it
             differs from line 1. If the student's label and the
             canonical concept are equivalent (e.g. both are
             "sample_size") we collapse to a single line so the cell
             stays compact for trivial mappings.
    """
    student_label = (student_label or "").strip()
    canonical = (canonical or "").strip()
    if not student_label:
        return canonical or "(unnamed)"
    if not canonical:
        return student_label
    if _is_equivalent_name(student_label, canonical):
        return student_label
    return f"{student_label}\n→ {canonical}"


def _is_equivalent_name(a: str, b: str) -> bool:
    """True when two field labels collapse to the same alphanumeric form."""
    norm_a = "".join(c for c in a.lower() if c.isalnum())
    norm_b = "".join(c for c in b.lower() if c.isalnum())
    return norm_a == norm_b and bool(norm_a)


def _render_extractor_reason(
    extractor_id: str,
    value: object,
    evidence: str,
    extractor_failed: bool,
) -> str:
    """Compose the per-extractor Reason cell.

    Three outcomes:
      * extractor_failed → ``"MODEL N/A"`` (a clear "did not produce
        a value" marker — never blamed on the student).
      * value present but no quote → ``"[llm:NAME] (no quote)"``.
      * value + quote → ``"[llm:NAME] <truncated quote>"``.
    """
    if extractor_failed:
        return "MODEL N/A"
    short = _short_extractor_name(extractor_id)
    if value is None or (isinstance(value, str) and not value.strip()):
        # Value missing but not flagged as a hard failure — usually
        # means the extractor explicitly said "couldn't find this".
        return f"[llm:{short}] (no value extracted)"
    if not evidence:
        return f"[llm:{short}] (no quote)"
    return _trunc(f"[llm:{short}] {evidence}", 200)


def _render_single(diff: FieldDiff, extractor_ids: list[str]) -> _RenderRow:
    """Render a non-grouped FieldDiff as a single table row."""
    field_label = _build_field_label(
        diff.student_raw_name or diff.field_name,
        diff.canonical_concept,
    )

    student_text = _fmt(diff.student_value)

    model_texts: list[str] = []
    model_reasons: list[str] = []
    # Pull per-extractor failed flags from the underlying ExtractedField
    # objects when available. FieldDiff stores values + evidence flat
    # lists in the same order as ``extractor_ids``; we don't have the
    # extractor_failed flag here directly, so we approximate "failed"
    # as "value is None AND evidence is empty" — extractor_failed
    # status is captured at the comparator level for MISSING_MODEL but
    # not propagated per-extractor on FieldDiff. Good enough for now.
    for j in range(len(extractor_ids)):
        value = diff.model_values[j] if j < len(diff.model_values) else None
        evid = diff.model_evidence[j] if j < len(diff.model_evidence) else ""
        model_texts.append(_fmt(value))
        # Heuristic: treat as failed when both value and evidence are
        # missing. This matches MISSING_MODEL extractor-gap semantics.
        treat_as_failed = (
            (value is None or (isinstance(value, str) and not value.strip()))
            and not evid
        )
        model_reasons.append(
            _render_extractor_reason(
                extractor_ids[j], value, evid, treat_as_failed
            )
        )

    return _RenderRow(
        field_label=field_label,
        student_text=student_text,
        model_texts=model_texts,
        status_value=getattr(diff.status, "value", str(diff.status)),
        model_reasons=model_reasons,
    )


def _render_merged(
    base: str,
    members: list[FieldDiff],
    extractor_ids: list[str],
) -> _RenderRow:
    """Combine T1DM + Control diffs into a single display row.

    The Student / AI cells contain two lines, each prefixed with the
    group label (e.g. ``T1DM: 12.90±1.30`` / ``Control: 12.96±1.12``).
    Reasons are also stacked per group, per extractor.
    """
    # Sort members in canonical group order (T1DM first, Control next).
    def _sort_key(d: FieldDiff) -> int:
        _, suf = _strip_group_suffix(d.field_name)
        try:
            return _GROUP_SUFFIXES.index(suf)
        except ValueError:
            return len(_GROUP_SUFFIXES)
    members = sorted(members, key=_sort_key)

    # Field label: pretty base name on line 1, canonical concept (if
    # any member declares one) on line 2 when it differs.
    student_raws = sorted({m.student_raw_name for m in members if m.student_raw_name})
    canonicals = sorted({
        m.canonical_concept for m in members
        if m.canonical_concept
        and not _is_equivalent_name(m.canonical_concept, base)
    })
    pretty_base = (
        " / ".join(student_raws)
        if student_raws
        else base.replace("_", " ").strip().title() or base
    )
    if canonicals:
        label = f"{pretty_base}\n→ {' / '.join(canonicals)}"
    else:
        label = pretty_base

    # Per-group lines for each cell.
    student_lines: list[str] = []
    model_lines_per_extractor: list[list[str]] = [[] for _ in extractor_ids]
    reason_lines_per_extractor: list[list[str]] = [[] for _ in extractor_ids]
    worst_status = "match"

    for m in members:
        _, suf = _strip_group_suffix(m.field_name)
        gl = _GROUP_LABEL.get(suf, suf.lstrip("_").upper() or "—")

        student_lines.append(f"{gl}: {_fmt(m.student_value)}")

        for j in range(len(extractor_ids)):
            value = m.model_values[j] if j < len(m.model_values) else None
            evid = m.model_evidence[j] if j < len(m.model_evidence) else ""
            model_lines_per_extractor[j].append(f"{gl}: {_fmt(value)}")

            treat_as_failed = (
                (value is None or (isinstance(value, str) and not value.strip()))
                and not evid
            )
            reason_text = _render_extractor_reason(
                extractor_ids[j], value, evid, treat_as_failed
            )
            reason_lines_per_extractor[j].append(f"{gl}: {reason_text}")

        # Roll up status using priority — worst wins.
        sv = getattr(m.status, "value", str(m.status))
        if _STATUS_PRIORITY.get(sv, 0) > _STATUS_PRIORITY.get(worst_status, 0):
            worst_status = sv

    return _RenderRow(
        field_label=label,
        student_text="\n".join(student_lines),
        model_texts=["\n".join(lines) for lines in model_lines_per_extractor],
        status_value=worst_status,
        model_reasons=[
            "\n".join(lines) for lines in reason_lines_per_extractor
        ],
    )

logger = structlog.get_logger(__name__)

# Colour constants (RGB tuples)
_GREEN = (198, 239, 206)
_RED = (255, 199, 206)
_YELLOW = (255, 235, 156)
_GREY = (217, 217, 217)
_WHITE = (255, 255, 255)
_DARK = (44, 62, 80)


def generate_docx_report(
    result: PipelineRunResult,
    output_path: Path | str,
) -> Path:
    """Generate a DOCX evaluation report from a pipeline result.

    Args:
        result: The completed pipeline result.
        output_path: Where to save the .docx file.

    Returns:
        The path to the saved report.
    """
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # -- Global style tweaks --
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ================================================================
    # COVER PAGE
    # ================================================================
    _add_cover_page(doc, result)

    # ================================================================
    # 1. EXECUTIVE SUMMARY
    # ================================================================
    doc.add_heading("1. Executive Summary", level=1)

    si = result.student_input
    summary_items = [
        f"Review: {si.review_title}",
        f"Student ID: {si.student_id}",
        f"Run ID: {result.run_id}",
        f"Papers analysed: {len(si.selected_papers)}",
        f"Completed: {result.completed_at or 'N/A'}",
    ]
    for item in summary_items:
        doc.add_paragraph(item, style="List Bullet")

    # Verdict
    if result.report:
        p = doc.add_paragraph()
        p.add_run("Verdict: ").bold = True
        p.add_run(result.report.summary)

    # ================================================================
    # 2. SEARCH VALIDATION (Step 1)
    # ================================================================
    doc.add_heading("2. Search Validation (Step 1)", level=1)

    # Disclaimer up front: Step 1 does NOT decide whether a paper exists;
    # that is Step 2's job. This avoids the "looks identical to Step 2"
    # confusion the report previously created.
    intro = doc.add_paragraph()
    intro_run = intro.add_run(
        "Note: Step 1 is informational only and does not contribute to the "
        "FAIL verdict. It checks search reproducibility — whether the "
        "student's reported per-database counts are consistent with what "
        "the same query returns today. Paper existence is the "
        "responsibility of Step 2 (Paper Verification)."
    )
    intro_run.italic = True

    # ----------------------------------------------------------------
    # 2.1 PRISMA-style multi-database identification count check
    # ----------------------------------------------------------------
    multi_db = result.multi_db_identification_check
    if multi_db is not None and multi_db.per_database:
        doc.add_heading(
            "2.1 Identification Count Check (multi-database)", level=2
        )
        doc.add_paragraph(
            "For each database the review claims to have searched, we "
            "either re-ran the query (PubMed / Europe PMC / OpenAlex) or "
            "marked the count UNVERIFIED when no free API is available "
            "(Embase, CINAHL, Cochrane, Web of Science)."
        )
        tbl = doc.add_table(
            rows=1 + len(multi_db.per_database), cols=5,
            style="Light Grid Accent 1",
        )
        for j, h in enumerate(
            ["Database", "Student", "AI", "Δ%", "Verdict"]
        ):
            _set_cell(tbl, 0, j, h, bold=True)
        verdict_colour = {
            "MATCH": _GREEN,
            "WARN": _YELLOW,
            "REFERENCE": _GREY,
            "UNVERIFIED": _GREY,
        }
        for i, row in enumerate(multi_db.per_database, 1):
            _set_cell(tbl, i, 0, row.database)
            _set_cell(
                tbl, i, 1,
                "—" if row.student_reported is None else str(row.student_reported),
            )
            _set_cell(
                tbl, i, 2,
                "—" if row.ai_reproduced is None else str(row.ai_reproduced),
            )
            _set_cell(
                tbl, i, 3,
                "—" if row.delta_pct is None else f"{row.delta_pct:.0f}%",
            )
            verdict_cell = tbl.cell(i, 4)
            verdict_cell.text = row.verdict
            _shade_cell(verdict_cell, verdict_colour.get(row.verdict, _WHITE))

        # Summary line for total reported / reproduced counts (when we
        # have anything meaningful to summarise).
        if (
            multi_db.student_reported_total is not None
            or multi_db.ai_total_unique is not None
        ):
            tot_p = doc.add_paragraph()
            tot_p.add_run("Totals: ").bold = True
            tot_p.add_run(
                f"student reported = "
                f"{multi_db.student_reported_total or '—'}, "
                f"AI reproduced (verifiable databases) = "
                f"{multi_db.ai_total_unique or '—'}."
            )
        if multi_db.note:
            note_p = doc.add_paragraph()
            note_run = note_p.add_run(multi_db.note)
            note_run.italic = True

    # ----------------------------------------------------------------
    # 2.2 PubMed reproducibility (legacy single-DB sub-check)
    # ----------------------------------------------------------------
    if result.search_result:
        doc.add_heading("2.2 Reproduced PubMed Query", level=2)
        sr = result.search_result
        tbl = doc.add_table(rows=4, cols=2, style="Light Grid Accent 1")
        _set_cell(tbl, 0, 0, "Search Query")
        _set_cell(tbl, 0, 1, sr.reconstructed_query[:120])
        _set_cell(tbl, 1, 0, "Reported Results")
        _set_cell(tbl, 1, 1, str(sr.reported_count))
        _set_cell(tbl, 2, 0, "Actual Results (PubMed)")
        _set_cell(tbl, 2, 1, str(sr.actual_count))
        _set_cell(tbl, 3, 0, "Reproducible?")
        _set_cell(tbl, 3, 1, "Yes" if sr.is_reproducible else "No")

    # NOTE: the former "2.3 Search Strategy Coverage of Selected Papers"
    # sub-section was removed on 2026-05-14 — it overlapped too heavily
    # with Step 2's paper verification and added little signal. The
    # underlying ``result.papers_in_search`` data is still computed by
    # the orchestrator (used for flags) but is no longer rendered here.

    # ================================================================
    # 3. PAPER VERIFICATION (Step 2)
    # ================================================================
    doc.add_heading("3. Paper Verification (Step 2)", level=1)

    if result.verification_results:
        tbl = doc.add_table(
            rows=1 + len(result.verification_results), cols=4,
            style="Light Grid Accent 1",
        )
        for j, h in enumerate(["Paper", "Status", "Confidence", "CrossRef Title"]):
            _set_cell(tbl, 0, j, h, bold=True)

        for i, vr in enumerate(result.verification_results, 1):
            _set_cell(tbl, i, 0, vr.reference.title[:55])
            status_cell = tbl.cell(i, 1)
            status_cell.text = vr.status.value
            colour = {"verified": _GREEN, "not_found": _RED}.get(
                vr.status.value, _YELLOW
            )
            _shade_cell(status_cell, colour)
            _set_cell(tbl, i, 2, f"{vr.confidence:.0%}")
            cr_title = (vr.matched_metadata or {}).get("title", "")
            _set_cell(tbl, i, 3, cr_title[:55])

    # ================================================================
    # 4. DATA EXTRACTION COMPARISON (Steps 3-4) — THE CORE TABLE
    # ================================================================
    doc.add_heading("4. Data Extraction Comparison (Steps 3-4)", level=1)

    doc.add_paragraph(
        "The table below compares the student's extracted data with "
        "AI-extracted values from each LLM, aligned on a canonical field "
        "schema (one row per logical field, after mapping student headers "
        "like 'N', 'EFT/EAT' to canonical names like 'sample_size', "
        "'eat_or_eft_t1dm'). Colours: Green = MATCH, Yellow = PARTIAL / "
        "NEEDS REVIEW, Red = DIFF, Grey = missing on one side."
    )

    # Identify distinct extractor names — prefer the authoritative list from
    # the pipeline run (populated even if some extractors failed); fall back to
    # scanning extracted_tables for backward compatibility.
    extractor_ids: list[str] = [
        e for e in result.extractor_ids if e != "student"
    ] if result.extractor_ids else []
    if not extractor_ids:
        for et in result.extracted_tables:
            if et.extractor_id not in extractor_ids and et.extractor_id != "student":
                extractor_ids.append(et.extractor_id)

    # Use the module-level _short_extractor_name helper so the table
    # header and the [llm:NAME] tags inside Reason cells stay in sync.
    _short_name = _short_extractor_name

    if result.comparison_results:
        for cr in result.comparison_results:
            # Match the comparison result back to the student's
            # ReferenceEntry so we can show authors / year / DOI / title.
            matched_ref = None
            for ref in result.student_input.selected_papers:
                norm_ref = (ref.doi or "").strip().lower()
                norm_cr = cr.paper_id.strip().lower()
                for prefix in ("https://doi.org/", "http://doi.org/", "doi:"):
                    norm_ref = norm_ref.removeprefix(prefix)
                    norm_cr = norm_cr.removeprefix(prefix)
                if norm_ref and norm_ref == norm_cr:
                    matched_ref = ref
                    break

            if matched_ref is not None:
                citation = _format_short_citation(matched_ref)
                paper_title = matched_ref.title or "(no title)"
                paper_doi = matched_ref.doi or "—"
            else:
                citation = (cr.paper_id[:40] or "(unidentified)").strip()
                paper_title = "(no title — could not match to selected_papers)"
                paper_doi = "—"

            # Heading line — keeps the document's outline pane tidy.
            doc.add_heading(f"Paper: {citation}", level=2)
            # Three labelled meta-lines (Title / Agreement / DOI), each
            # rendered with the label in bold and the value in plain
            # weight. Format finalised with the user on 2026-05-10.
            _add_meta_line(doc, "Title", paper_title)
            _add_meta_line(doc, "Agreement", f"{cr.agreement_rate:.0%}")
            _add_meta_line(doc, "DOI", paper_doi)

            # New column layout (2026-05-10):
            #   Field | Student | <LLM-A val> | <LLM-B val> | Status |
            #   Reason (LLM-A) | Reason (LLM-B)
            #
            # Each extractor gets its own Reason column (placed after
            # Status) so readers can see why each model arrived at its
            # value side-by-side. Failed extractors render as "MODEL N/A"
            # in their Reason cell.
            col_headers = ["Field", "Student"]
            for eid in extractor_ids:
                col_headers.append(_short_name(eid))
            col_headers.append("Status")
            for eid in extractor_ids:
                col_headers.append(f"Reason ({_short_name(eid)})")
            n_cols = len(col_headers)
            n_extractors = len(extractor_ids)
            status_col_idx = 2 + n_extractors  # position of the Status column
            first_reason_col = status_col_idx + 1

            if not cr.field_diffs:
                doc.add_paragraph("(No fields to compare for this paper.)")
                doc.add_paragraph()
                continue

            # Pre-process: collapse (base, t1dm)+(base, control) pairs into
            # single merged rows so the report's row layout matches the
            # original "Characteristics of included studies" table layout
            # (one row per field, with both group values stacked inside).
            display_rows = _merge_grouped_diffs(cr.field_diffs, extractor_ids)

            tbl = doc.add_table(
                rows=1 + len(display_rows), cols=n_cols,
                style="Light Grid Accent 1",
            )
            for j, h in enumerate(col_headers):
                _set_cell(tbl, 0, j, h, bold=True)

            for i, row in enumerate(display_rows, 1):
                _set_cell(tbl, i, 0, row.field_label)
                _set_cell(tbl, i, 1, row.student_text)
                for j in range(n_extractors):
                    cell_text = row.model_texts[j] if j < len(row.model_texts) else "—"
                    tbl.cell(i, 2 + j).text = cell_text

                # Status column — colour-coded by the rolled-up status.
                status_cell = tbl.cell(i, status_col_idx)
                sv = row.status_value
                if sv == "match":
                    status_cell.text = "✓ MATCH"
                    _shade_cell(status_cell, _GREEN)
                elif sv == "partial_match":
                    status_cell.text = "~ PARTIAL"
                    _shade_cell(status_cell, _YELLOW)
                elif sv == "missing_model":
                    status_cell.text = "MODEL N/A"
                    _shade_cell(status_cell, _GREY)
                elif sv == "missing_student":
                    status_cell.text = "STUDENT N/A"
                    _shade_cell(status_cell, _GREY)
                elif sv == "not_comparable":
                    status_cell.text = "N/A"
                    _shade_cell(status_cell, _GREY)
                elif sv == "needs_review":
                    status_cell.text = "? REVIEW"
                    _shade_cell(status_cell, _YELLOW)
                else:  # diff
                    status_cell.text = "✗ DIFF"
                    _shade_cell(status_cell, _RED)

                # Per-extractor Reason cells.
                for j in range(n_extractors):
                    reason_text = (
                        row.model_reasons[j]
                        if j < len(row.model_reasons)
                        else "—"
                    )
                    reason_cell = tbl.cell(i, first_reason_col + j)
                    _set_cell(tbl, i, first_reason_col + j, reason_text)
                    if reason_text == "MODEL N/A":
                        _shade_cell(reason_cell, _GREY)

            # Three blank paragraphs as a visual separator between papers
            # (per user request 2026-05-10) — gives space when printing.
            for _ in range(3):
                doc.add_paragraph()

    # ================================================================
    # 5. CHARTS
    # ================================================================
    doc.add_heading("5. Visualisation", level=1)

    # 5a. Agreement rate per paper (bar chart)
    if result.comparison_results:
        doc.add_heading("Agreement Rate per Paper", level=2)
        # Build a paper_id → "Surname Year" label map so the chart shows
        # human-readable citations rather than raw DOIs.
        paper_id_to_citation: dict[str, str] = {}
        for ref in result.student_input.selected_papers:
            doi_key = (ref.doi or "").strip().lower()
            if not doi_key:
                continue
            for prefix in ("https://doi.org/", "http://doi.org/", "doi:"):
                doi_key = doi_key.removeprefix(prefix)
            authors = list(getattr(ref, "authors", None) or [])
            surname = _surname(authors[0]) if authors else ""
            year = ref.year
            if surname and year:
                label = f"{surname} {year}"
            elif surname:
                label = surname
            elif year:
                label = str(year)
            else:
                label = doi_key.split("/")[-1][:25]
            paper_id_to_citation[doi_key] = label
        chart_bytes = _chart_agreement_per_paper(
            result.comparison_results,
            paper_id_to_citation,
        )
        if chart_bytes:
            doc.add_picture(io.BytesIO(chart_bytes), width=Inches(5.8))

    # 5b. Agreement rate per field (averaged across papers)
    if result.comparison_results:
        doc.add_heading("Agreement Rate per Field", level=2)
        chart_bytes = _chart_agreement_per_field(result.comparison_results)
        if chart_bytes:
            doc.add_picture(io.BytesIO(chart_bytes), width=Inches(5.8))

    # ================================================================
    # 6. FLAGS APPENDIX
    # ================================================================
    doc.add_heading("6. All Flags", level=1)

    if result.all_flags:
        tbl = doc.add_table(
            rows=1 + len(result.all_flags), cols=4,
            style="Light Grid Accent 1",
        )
        for j, h in enumerate(["Severity", "Step", "Code", "Message"]):
            _set_cell(tbl, 0, j, h, bold=True)
        for i, flag in enumerate(result.all_flags, 1):
            sev_cell = tbl.cell(i, 0)
            sev_cell.text = flag.severity.value.upper()
            colour = {"error": _RED, "warning": _YELLOW, "info": _GREEN}.get(
                flag.severity.value, _WHITE
            )
            _shade_cell(sev_cell, colour)
            _set_cell(tbl, i, 1, flag.step.value)
            _set_cell(tbl, i, 2, flag.code)
            _set_cell(tbl, i, 3, flag.message[:100])
    else:
        doc.add_paragraph("No flags raised.")

    # ================================================================
    # SAVE
    # ================================================================
    doc.save(str(output_path))
    logger.info("docx_report_saved", path=str(output_path))
    return output_path


# ======================================================================
# Chart helpers (matplotlib)
# ======================================================================


def _chart_agreement_per_paper(
    comparisons: list[TableComparisonResult],
    paper_id_to_citation: dict[str, str] | None = None,
) -> bytes | None:
    """Generate a horizontal bar chart: agreement rate per paper.

    Args:
        comparisons: Per-paper comparison results.
        paper_id_to_citation: Optional map from normalised DOI → human
            citation (e.g. ``"Borghaei 2015"``). When supplied, the chart
            uses these labels; otherwise it falls back to a short DOI
            suffix.
    """
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except ImportError:
        logger.warning("matplotlib_not_installed")
        return None

    citations = paper_id_to_citation or {}

    labels = []
    rates = []
    colours = []
    for cr in comparisons:
        # Prefer the human citation when we have one for this paper_id.
        norm_id = cr.paper_id.strip().lower()
        for prefix in ("https://doi.org/", "http://doi.org/", "doi:"):
            norm_id = norm_id.removeprefix(prefix)
        short = citations.get(norm_id) or cr.paper_id
        if "/" in short and short == cr.paper_id:
            short = short.split("/")[-1]
        if len(short) > 30:
            short = short[:27] + "..."
        labels.append(short)
        rates.append(cr.agreement_rate * 100)
        if cr.agreement_rate >= 0.7:
            colours.append("#27ae60")
        elif cr.agreement_rate >= 0.4:
            colours.append("#f39c12")
        else:
            colours.append("#e74c3c")

    fig, ax = plt.subplots(figsize=(8, max(3, len(labels) * 0.5)))
    y_pos = range(len(labels))
    bars = ax.barh(y_pos, rates, color=colours, edgecolor="white", height=0.6)
    ax.set_yticks(y_pos)
    ax.set_yticklabels(labels, fontsize=8)
    ax.set_xlabel("Agreement Rate (%)")
    ax.set_xlim(0, 105)
    ax.set_title("Student vs AI: Agreement Rate per Paper")
    ax.axvline(x=70, color="#27ae60", linestyle="--", alpha=0.5, label="70% threshold")
    ax.legend(fontsize=7)

    # Add percentage labels on bars
    for bar, rate in zip(bars, rates):
        ax.text(
            bar.get_width() + 1, bar.get_y() + bar.get_height() / 2,
            f"{rate:.0f}%", va="center", fontsize=8,
        )

    ax.invert_yaxis()
    plt.tight_layout()

    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    return buf.getvalue()


def _chart_agreement_per_field(
    comparisons: list[TableComparisonResult],
) -> bytes | None:
    """Generate a vertical bar chart: agreement rate per field (averaged)."""
    try:
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except ImportError:
        return None

    # Aggregate: for each field_name, count match / total across papers
    field_stats: dict[str, list[bool]] = {}
    for cr in comparisons:
        for diff in cr.field_diffs:
            field_stats.setdefault(diff.field_name, []).append(diff.is_consistent)

    if not field_stats:
        return None

    labels = list(field_stats.keys())
    rates = [
        sum(v) / len(v) * 100 if v else 0
        for v in field_stats.values()
    ]
    colours = [
        "#27ae60" if r >= 70 else "#f39c12" if r >= 40 else "#e74c3c"
        for r in rates
    ]

    fig, ax = plt.subplots(figsize=(max(6, len(labels) * 0.6), 4))
    x_pos = range(len(labels))
    bars = ax.bar(x_pos, rates, color=colours, edgecolor="white", width=0.6)
    ax.set_xticks(x_pos)
    ax.set_xticklabels(labels, rotation=45, ha="right", fontsize=8)
    ax.set_ylabel("Agreement Rate (%)")
    ax.set_ylim(0, 110)
    ax.set_title("Student vs AI: Agreement Rate per Field")
    ax.axhline(y=70, color="#27ae60", linestyle="--", alpha=0.5)

    for bar, rate in zip(bars, rates):
        ax.text(
            bar.get_x() + bar.get_width() / 2, bar.get_height() + 2,
            f"{rate:.0f}%", ha="center", fontsize=7,
        )

    plt.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=150)
    plt.close(fig)
    return buf.getvalue()


# ======================================================================
# Document helpers
# ======================================================================


def _add_cover_page(doc, result: PipelineRunResult) -> None:
    """Add a professional cover page."""
    from docx.shared import Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    # Spacer
    for _ in range(4):
        doc.add_paragraph()

    # Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Literature Integrity Detection Report")
    run.font.size = Pt(26)
    run.bold = True

    # Subtitle
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run(result.student_input.review_title)
    run.font.size = Pt(14)
    run.italic = True

    doc.add_paragraph()

    # Metadata
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Student: {result.student_input.student_id}\n").font.size = Pt(11)
    meta.add_run(f"Run ID: {result.run_id}\n").font.size = Pt(11)
    meta.add_run(
        f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
    ).font.size = Pt(11)

    # Extractors used — prefer authoritative list from pipeline run
    eids = list(result.extractor_ids) if result.extractor_ids else []
    if not eids:
        for et in result.extracted_tables:
            if et.extractor_id not in eids:
                eids.append(et.extractor_id)
    if eids:
        meta.add_run(f"LLM Extractors: {', '.join(eids)}\n").font.size = Pt(11)

    doc.add_page_break()


def _set_cell(
    table, row: int, col: int, text: str, bold: bool = False
) -> None:
    """Set text in a table cell."""
    from docx.shared import Pt

    cell = table.cell(row, col)
    cell.text = text
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(8)
            if bold:
                run.bold = True


def _shade_cell(cell, rgb: tuple[int, int, int]) -> None:
    """Apply a background colour to a table cell."""
    from docx.oxml.ns import qn

    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.find(qn("w:shd"))
    if shading_elm is None:
        from docx.oxml import OxmlElement
        shading_elm = OxmlElement("w:shd")
        shading.append(shading_elm)
    hex_colour = f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"
    shading_elm.set(qn("w:fill"), hex_colour)
    shading_elm.set(qn("w:val"), "clear")


def _fmt(value: object) -> str:
    """Format a value for display in a table cell."""
    if value is None:
        return "N/A"
    s = str(value)
    if len(s) > 35:
        s = s[:32] + "..."
    return s


def _trunc(text: str, limit: int) -> str:
    """Hard-truncate with ellipsis for evidence cells."""
    if not text:
        return ""
    if len(text) <= limit:
        return text
    return text[: max(0, limit - 3)] + "..."


# ----------------------------------------------------------------------
# Citation / labelled-line helpers (used by the per-paper headings)
# ----------------------------------------------------------------------


def _surname(author: str) -> str:
    """Strip trailing initials from an author string.

    "Borghaei H"            → "Borghaei"
    "Reck M"                → "Reck"
    "de Gonzalo-Calvo D"    → "de Gonzalo-Calvo"  (preserves particles)
    "Smith"                 → "Smith"
    """
    tokens = (author or "").strip().split()
    while tokens and re.fullmatch(r"[A-Z][A-Za-z]?\.?", tokens[-1]):
        tokens.pop()
    return " ".join(tokens) if tokens else (author or "").strip()


def _format_short_citation(ref) -> str:
    """Build a short citation like ``Borghaei et al. (2015)``.

    * Multiple authors  → "{first_surname} et al. ({year})"
    * Single author     → "{surname} ({year})"
    * Year only         → "(year)"
    * Nothing usable    → "Unknown"
    """
    authors = list(getattr(ref, "authors", None) or [])
    year = getattr(ref, "year", None)
    surname = _surname(authors[0]) if authors else ""
    year_str = f"({year})" if year else ""

    if surname and len(authors) > 1:
        return (f"{surname} et al. {year_str}").strip()
    if surname:
        return (f"{surname} {year_str}").strip()
    return year_str or "Unknown"


def _add_meta_line(doc, label: str, value: str) -> None:
    """Add a paragraph rendered as ``**Label:** value``.

    Used by the per-paper Step 4 sections so each paper carries a
    consistent four-line header (Paper / Title / Agreement / DOI) with
    bold labels and plain-weight values.
    """
    from docx.shared import Pt

    p = doc.add_paragraph()
    label_run = p.add_run(f"{label}: ")
    label_run.bold = True
    label_run.font.size = Pt(10)
    value_run = p.add_run(value or "—")
    value_run.font.size = Pt(10)
